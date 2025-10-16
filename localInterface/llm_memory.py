# llm_memory.py
import os
import shutil
import pandas as pd
from langchain.schema import Document
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredPowerPointLoader,
    CSVLoader,
)
from docx import Document
from langchain.schema import Document as LangchainDoc
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain_community.vectorstores import FAISS

# Absolute paths (important)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_PATH = os.path.join(BASE_DIR, "localInterface", "data")
DB_FAISS_PATH = os.path.join(BASE_DIR, "vectorstore", "db_faiss")

# Config: how we convert excel rows -> documents
EXCEL_ROW_TO_TEXT_JOINER = " | "  # change if you want different formatting

def extract_text_from_excel(file_path):
    excel_data = pd.read_excel(file_path, sheet_name=None)
    text_data = ""
    for sheet_name, df in excel_data.items():
        text_data += f"\nSheet: {sheet_name}\n"
        df = df.fillna("")
        for _, row in df.iterrows():
            row_text = ", ".join([
                f"{col.strip()}: {str(val).strip()}"
                for col, val in row.items() if str(val).strip() != ""
            ])
            text_data += row_text + "\n"
    return text_data.strip()

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text_data = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text_data += para.text.strip() + "\n"
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                text_data += ", ".join(cells) + "\n"
    return text_data.strip()

def load_documents():
    """Load supported documents from DATA_PATH and return list[Document]."""
    documents = []
    os.makedirs(DATA_PATH, exist_ok=True)

    files = sorted(os.listdir(DATA_PATH))
    print("Found files in data/:", files)

    for file in files:
        path = os.path.join(DATA_PATH, file)
        if not os.path.isfile(path):
            continue

        ext = os.path.splitext(file)[1].lower().lstrip(".")
        try:
            if ext == "pdf":
                print("Loading PDF:", file)
                loader = PyPDFLoader(path)
                docs = loader.load()
                for d in docs:
                    d.metadata = d.metadata or {}
                    d.metadata.setdefault("source", file)
                documents.extend(docs)

            elif ext in ("doc", "docx"):
                print("Loading Word doc:", file)
                content = extract_text_from_docx(path)
                documents.append(LangchainDoc(page_content=content, metadata={"source": file}))

            elif ext in ("ppt", "pptx"):
                print("Loading PPT:", file)
                loader = UnstructuredPowerPointLoader(path)
                docs = loader.load()
                for d in docs:
                    d.metadata = d.metadata or {}
                    d.metadata.setdefault("source", file)
                documents.extend(docs)

            elif ext == "csv":
                print("Loading CSV:", file)
                loader = CSVLoader(path)
                docs = loader.load()
                for d in docs:
                    d.metadata = d.metadata or {}
                    d.metadata.setdefault("source", file)
                documents.extend(docs)

            elif ext in ("xls", "xlsx"):
                print("Loading Excel:", file)
                content = extract_text_from_excel(path)
                documents.append(LangchainDoc(page_content=content, metadata={"source": file}))

            elif ext == "txt":
                print("Loading TXT:", file)
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read().strip()
                    if content:
                        documents.append(Document(page_content=content, metadata={"source": file}))

            else:
                print("Skipping unsupported file type:", file)

        except Exception as e:
            print(f"⚠️ Error loading {file}: {e}")

    return documents


def rebuild_database():
    """Full rebuild: clear old DB, load documents, split, embed, save FAISS."""
    # ensure data folder exists
    os.makedirs(DATA_PATH, exist_ok=True)

    documents = load_documents()
    if not documents:
        # clear DB when nothing to index to avoid stale answers
        if os.path.exists(DB_FAISS_PATH):
            shutil.rmtree(DB_FAISS_PATH)
            print("Cleared old FAISS DB at", DB_FAISS_PATH)
        raise ValueError("No documents found in data/")

    print(f"Loaded {len(documents)} raw documents. Splitting into chunks...")
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=80)
    text_chunks = splitter.split_documents(documents)
    print(f"Total chunks created: {len(text_chunks)}")

    # normalize metadata
    for doc in text_chunks:
        doc.metadata = doc.metadata or {}
        src = doc.metadata.get("source", "unknown")
        page = doc.metadata.get("page", doc.metadata.get("row", "unknown"))
        doc.metadata["source"] = f"{src} - page {page}"

    # Embedding model (CPU by default). If you want GPU set {"device":"cuda"}.
    embedding_model = SentenceTransformerEmbeddings(
        model_name="intfloat/e5-small-v2",
        model_kwargs={"device": "cpu"}
    )

    # build and save FAISS
    db = FAISS.from_documents(text_chunks, embedding_model)
    os.makedirs(os.path.dirname(DB_FAISS_PATH), exist_ok=True)
    db.save_local(DB_FAISS_PATH)
    print("FAISS DB saved to:", DB_FAISS_PATH)
    return True


def get_vectorstore():
    """Load FAISS DB (used at runtime)."""
    if not os.path.exists(DB_FAISS_PATH):
        raise ValueError("FAISS DB not found. Rebuild first.")
    embedding_model = SentenceTransformerEmbeddings(
        model_name="intfloat/e5-small-v2",
        model_kwargs={"device": "cpu"}
    )
    db = FAISS.load_local(DB_FAISS_PATH, embedding_model, allow_dangerous_deserialization=True)
    return db
